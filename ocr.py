from __future__ import annotations

import argparse
import json
import os
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any, Iterable

import numpy as np
from paddleocr import PaddleOCR
from pdf2image import convert_from_path
from PIL import ImageGrab


IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".bmp", ".tif", ".tiff", ".webp"}
DEFAULT_OUTPUT_DIR = "output"

_OCR_ENGINE: PaddleOCR | None = None


@dataclass
class OCRLine:
	text: str
	confidence: float


@dataclass
class OCRPage:
	page_number: int
	lines: list[OCRLine]
	text: str


@dataclass
class OCRDocument:
	source_file: str
	file_type: str
	pages: list[OCRPage]
	full_text: str


def get_ocr_engine() -> PaddleOCR:
	global _OCR_ENGINE
	if _OCR_ENGINE is None:
		_OCR_ENGINE = PaddleOCR(use_angle_cls=True, lang="en", show_log=False)
	return _OCR_ENGINE


def is_supported_image(file_path: Path) -> bool:
	return file_path.suffix.lower() in IMAGE_EXTENSIONS


def is_pdf(file_path: Path) -> bool:
	return file_path.suffix.lower() == ".pdf"


def _is_recognition_line(item: Any) -> bool:
	return (
		isinstance(item, (list, tuple))
		and len(item) == 2
		and isinstance(item[1], (list, tuple))
		and len(item[1]) >= 2
		and isinstance(item[1][0], str)
	)


def _walk_ocr_result(result: Any) -> Iterable[Any]:
	if result is None:
		return

	if _is_recognition_line(result):
		yield result
		return

	if isinstance(result, (list, tuple)):
		for item in result:
			if _is_recognition_line(item):
				yield item
			elif isinstance(item, (list, tuple)):
				yield from _walk_ocr_result(item)


def _normalize_ocr_result(result: Any) -> list[OCRLine]:
	lines: list[OCRLine] = []

	for item in _walk_ocr_result(result):
		text = str(item[1][0]).strip()
		confidence = float(item[1][1])

		if text:
			lines.append(OCRLine(text=text, confidence=confidence))

	return lines


def get_image_from_clipboard() -> str | None:
	"""
	Get image from clipboard and save it to output directory.
	Returns the path to the saved image, or None if no image found.
	"""
	try:
		image = ImageGrab.grabclipboard()
		if image is None:
			print("Error: No image found in clipboard!")
			return None

		# Create output directory if it doesn't exist
		os.makedirs(DEFAULT_OUTPUT_DIR, exist_ok=True)

		# Save clipboard image with timestamp
		import time
		timestamp = int(time.time() * 1000)
		clipboard_image_path = os.path.join(
			DEFAULT_OUTPUT_DIR,
			f"clipboard_{timestamp}.png"
		)

		image.save(clipboard_image_path)
		print(f"Image from clipboard saved to: {clipboard_image_path}")
		return clipboard_image_path

	except Exception as e:
		print(f"Error accessing clipboard: {e}")
		return None


def extract_text_from_image(image_path: str | Path) -> OCRPage:
	engine = get_ocr_engine()
	raw_result = engine.ocr(str(image_path), cls=True)
	lines = _normalize_ocr_result(raw_result)
	page_text = "\n".join(line.text for line in lines)
	return OCRPage(page_number=1, lines=lines, text=page_text)


def extract_text_from_pdf(pdf_path: str | Path) -> list[OCRPage]:
	engine = get_ocr_engine()
	pages: list[OCRPage] = []

	for page_number, page_image in enumerate(convert_from_path(str(pdf_path)), start=1):
		raw_result = engine.ocr(np.array(page_image.convert("RGB")), cls=True)
		lines = _normalize_ocr_result(raw_result)
		page_text = "\n".join(line.text for line in lines)
		pages.append(OCRPage(page_number=page_number, lines=lines, text=page_text))

	return pages


def build_document(file_path: str | Path) -> OCRDocument:
	path = Path(file_path)

	if is_supported_image(path):
		pages = [extract_text_from_image(path)]
		file_type = "image"
	elif is_pdf(path):
		pages = extract_text_from_pdf(path)
		file_type = "pdf"
	else:
		raise ValueError(f"Unsupported file format: {path.suffix}")

	if len(pages) == 1:
		full_text = pages[0].text
	else:
		sections: list[str] = []
		for page in pages:
			sections.append(f"========== PAGE {page.page_number} ==========")
			if page.text:
				sections.append(page.text)
		full_text = "\n".join(sections).strip()

	return OCRDocument(
		source_file=str(path),
		file_type=file_type,
		pages=pages,
		full_text=full_text,
	)


def save_txt(document: OCRDocument, input_file: str | Path, output_dir: str | Path) -> Path:
	output_path = Path(output_dir)
	output_path.mkdir(parents=True, exist_ok=True)

	output_file = output_path / f"{Path(input_file).stem}.txt"
	output_file.write_text(document.full_text, encoding="utf-8")
	return output_file


def save_json(document: OCRDocument, input_file: str | Path, output_dir: str | Path) -> Path:
	output_path = Path(output_dir)
	output_path.mkdir(parents=True, exist_ok=True)

	output_file = output_path / f"{Path(input_file).stem}.json"
	payload = {
		"source_file": document.source_file,
		"file_type": document.file_type,
		"page_count": len(document.pages),
		"pages": [
			{
				"page_number": page.page_number,
				"text": page.text,
				"lines": [asdict(line) for line in page.lines],
			}
			for page in document.pages
		],
		"full_text": document.full_text,
	}

	output_file.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")
	return output_file


def save_docx(document: OCRDocument, input_file: str | Path, output_dir: str | Path) -> Path:
	try:
		from docx import Document
	except ImportError as exc:
		raise RuntimeError(
			"DOCX output requires the python-docx package. Install it and try again."
		) from exc

	output_path = Path(output_dir)
	output_path.mkdir(parents=True, exist_ok=True)

	output_file = output_path / f"{Path(input_file).stem}.docx"
	document_writer = Document()
	document_writer.add_heading("OCR Extracted Text", level=1)
	document_writer.add_paragraph(f"Source file: {document.source_file}")
	document_writer.add_paragraph(f"File type: {document.file_type}")

	for page in document.pages:
		document_writer.add_heading(f"Page {page.page_number}", level=2)
		if page.text:
			for line in page.lines:
				document_writer.add_paragraph(line.text)
		else:
			document_writer.add_paragraph("No text detected.")

	document_writer.save(str(output_file))
	return output_file


def create_parser() -> argparse.ArgumentParser:
	parser = argparse.ArgumentParser(
		description="Extract text from images or scanned PDFs using PaddleOCR."
	)
	parser.add_argument(
		"input_path",
		nargs="?",
		help="Path to an image file or scanned PDF document.",
	)
	parser.add_argument(
		"--output-dir",
		default=DEFAULT_OUTPUT_DIR,
		help="Directory where extracted files will be written.",
	)
	parser.add_argument(
		"--docx",
		action="store_true",
		help="Also write the result to a DOCX file.",
	)
	parser.add_argument(
		"--json",
		action="store_true",
		help="Also write the result to a JSON file.",
	)
	return parser


def run(file_path: str, output_dir: str = DEFAULT_OUTPUT_DIR, export_docx: bool = False, export_json: bool = False) -> list[Path]:
	path = Path(file_path)

	if not path.exists():
		raise FileNotFoundError(f"File not found: {path}")

	document = build_document(path)
	written_files = [save_txt(document, path, output_dir)]

	if export_docx:
		written_files.append(save_docx(document, path, output_dir))

	if export_json:
		written_files.append(save_json(document, path, output_dir))

	return written_files


def main(argv: list[str] | None = None) -> int:
	parser = create_parser()
	args = parser.parse_args(argv)

	input_path = args.input_path
	if not input_path:
		print("Enter image or PDF path (or type 'clipboard' to paste from clipboard): ", end="")
		input_path = input().strip()
	
	# Handle clipboard input
	if input_path.lower() == "clipboard":
		input_path = get_image_from_clipboard()
		if input_path is None:
			return 1
	
	if not input_path:
		print("Error: no input file provided.")
		return 1

	try:
		written_files = run(
			input_path,
			output_dir=args.output_dir,
			export_docx=args.docx,
			export_json=args.json,
		)
	except Exception as exc:
		print(f"Error: {exc}")
		return 1

	print("Extracted text saved to:")
	for output_file in written_files:
		print(f"- {output_file}")

	return 0


if __name__ == "__main__":
	raise SystemExit(main())
